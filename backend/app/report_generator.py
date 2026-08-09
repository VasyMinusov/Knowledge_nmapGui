# backend/app/report_generator.py
import io
import os
from datetime import datetime
from typing import Dict, List, Any
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from docx import Document
from .nmap_wrapper import parse_nmap_xml
from .database import get_scan_by_id
from .models import HostInfo  # <-- импорт модели

def generate_html_report(scan_data: Dict, hosts: List[HostInfo]) -> str:
    """Генерирует HTML-отчёт."""
    html = f"""
    <!DOCTYPE html>
    <html>
    <head><meta charset="utf-8"><title>Nmap Scan Report</title>
    <style>
        body {{ font-family: monospace; background: #0c120c; color: #ccff66; padding: 20px; }}
        h1 {{ color: #66ff33; text-shadow: 0 0 10px #66ff33; }}
        .host {{ border: 2px solid #66ff33; margin: 10px 0; padding: 10px; }}
        .port {{ display: inline-block; border: 1px solid #66ff33; padding: 2px 8px; margin: 4px; }}
        .open {{ color: #66ff33; }}
        .closed {{ color: #777; }}
        table {{ border-collapse: collapse; width: 100%; }}
        th, td {{ border: 1px solid #66ff33; padding: 6px; text-align: left; }}
        th {{ background: #1f301f; }}
    </style>
    </head>
    <body>
    <h1>Nmap Scan Report</h1>
    <p><strong>Scan ID:</strong> {scan_data['scan_id']}</p>
    <p><strong>Targets:</strong> {scan_data['targets']}</p>
    <p><strong>Profile:</strong> {scan_data['profile']}</p>
    <p><strong>Start:</strong> {scan_data['start_time']}</p>
    <p><strong>End:</strong> {scan_data.get('end_time', 'N/A')}</p>
    <p><strong>Status:</strong> {scan_data['status']}</p>
    <p><strong>Summary:</strong> {scan_data.get('summary', '')}</p>
    <hr>
    """
    for host in hosts:
        html += f"""
        <div class="host">
            <h2>{host.ip} ({host.hostname or 'unknown'}) – {host.status}</h2>
            <p>OS: {host.os or 'N/A'} | Uptime: {host.uptime or 'N/A'}s</p>
            <h3>Open Ports</h3>
            <table>
                <tr><th>Port</th><th>Protocol</th><th>Service</th><th>Version</th></tr>
        """
        for p in host.ports:
            if p.state == 'open':
                html += f"""
                <tr>
                    <td>{p.port}</td>
                    <td>{p.protocol}</td>
                    <td>{p.service or '—'}</td>
                    <td>{p.version or '—'}</td>
                </tr>
                """
        html += "</table></div>"
    html += "</body></html>"
    return html

def generate_docx_report(scan_data: Dict, hosts: List[HostInfo]) -> io.BytesIO:
    """Генерирует DOCX-отчёт и возвращает BytesIO."""
    doc = Document()
    doc.add_heading('Nmap Scan Report', 0)
    doc.add_paragraph(f"Scan ID: {scan_data['scan_id']}")
    doc.add_paragraph(f"Targets: {scan_data['targets']}")
    doc.add_paragraph(f"Profile: {scan_data['profile']}")
    doc.add_paragraph(f"Start: {scan_data['start_time']}")
    doc.add_paragraph(f"End: {scan_data.get('end_time', 'N/A')}")
    doc.add_paragraph(f"Status: {scan_data['status']}")
    doc.add_paragraph(f"Summary: {scan_data.get('summary', '')}")
    doc.add_paragraph('\n')

    for host in hosts:
        doc.add_heading(f"{host.ip} ({host.hostname or 'unknown'})", level=1)
        doc.add_paragraph(f"Status: {host.status} | OS: {host.os or 'N/A'} | Uptime: {host.uptime or 'N/A'}s")
        doc.add_heading('Open Ports', level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Port'
        hdr_cells[1].text = 'Protocol'
        hdr_cells[2].text = 'Service'
        hdr_cells[3].text = 'Version'
        for p in host.ports:
            if p.state == 'open':
                row_cells = table.add_row().cells
                row_cells[0].text = str(p.port)
                row_cells[1].text = p.protocol
                row_cells[2].text = p.service or '—'
                row_cells[3].text = p.version or '—'
        doc.add_paragraph('\n')
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_pdf_report(scan_data: Dict, hosts: List[HostInfo]) -> io.BytesIO:
    """Генерирует PDF-отчёт с помощью reportlab."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
    styles = getSampleStyleSheet()
    title_style = styles['Title']
    heading_style = styles['Heading1']
    normal_style = styles['Normal']

    mono_style = ParagraphStyle('Mono', parent=normal_style, fontName='Courier', fontSize=10)

    story = []

    story.append(Paragraph("Nmap Scan Report", title_style))
    story.append(Spacer(1, 0.2*inch))
    story.append(Paragraph(f"<b>Scan ID:</b> {scan_data['scan_id']}", mono_style))
    story.append(Paragraph(f"<b>Targets:</b> {scan_data['targets']}", mono_style))
    story.append(Paragraph(f"<b>Profile:</b> {scan_data['profile']}", mono_style))
    story.append(Paragraph(f"<b>Start:</b> {scan_data['start_time']}", mono_style))
    story.append(Paragraph(f"<b>End:</b> {scan_data.get('end_time', 'N/A')}", mono_style))
    story.append(Paragraph(f"<b>Status:</b> {scan_data['status']}", mono_style))
    story.append(Paragraph(f"<b>Summary:</b> {scan_data.get('summary', '')}", mono_style))
    story.append(Spacer(1, 0.3*inch))

    for host in hosts:
        story.append(Paragraph(f"{host.ip} ({host.hostname or 'unknown'})", heading_style))
        story.append(Paragraph(f"Status: {host.status} | OS: {host.os or 'N/A'} | Uptime: {host.uptime or 'N/A'}s", normal_style))
        story.append(Paragraph("Open Ports:", normal_style))
        data = [['Port', 'Protocol', 'Service', 'Version']]
        for p in host.ports:
            if p.state == 'open':
                data.append([
                    str(p.port),
                    p.protocol,
                    p.service or '—',
                    p.version or '—'
                ])
        if len(data) > 1:
            table = Table(data, colWidths=[0.8*inch, 0.8*inch, 1.5*inch, 1.5*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.grey),
                ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('FONTSIZE', (0,0), (-1,0), 12),
                ('BOTTOMPADDING', (0,0), (-1,0), 8),
                ('BACKGROUND', (0,1), (-1,-1), colors.beige),
                ('GRID', (0,0), (-1,-1), 1, colors.black),
            ]))
            story.append(table)
        else:
            story.append(Paragraph("No open ports found.", normal_style))
        story.append(Spacer(1, 0.2*inch))

    doc.build(story)
    buffer.seek(0)
    return buffer

def generate_report(scan_id: str, format: str = 'html'):
    """Основная функция генерации отчёта. Возвращает кортеж (filename, content, media_type)."""
    scan_data = get_scan_by_id(scan_id)
    if not scan_data:
        return None, None, None

    result_path = scan_data.get('result_path')
    hosts = []
    if result_path and os.path.exists(result_path):
        with open(result_path, 'r') as f:
            xml_data = f.read()
        parsed = parse_nmap_xml(xml_data)
        hosts = parsed.get('hosts', [])  # уже список HostInfo

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"scan_{scan_id[:8]}_{timestamp}"

    if format == 'html':
        content = generate_html_report(scan_data, hosts)
        return f"{filename}.html", content, "text/html"
    elif format == 'docx':
        buffer = generate_docx_report(scan_data, hosts)
        return f"{filename}.docx", buffer.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif format == 'pdf':
        buffer = generate_pdf_report(scan_data, hosts)
        return f"{filename}.pdf", buffer.getvalue(), "application/pdf"
    else:
        return None, None, None